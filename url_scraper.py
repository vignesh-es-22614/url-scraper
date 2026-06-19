"""
URL Scraper Tool
----------------
Reads URLs from an Excel (.xlsx) or CSV file, scrapes text content
from each URL, and saves results into a formatted Word document.

Usage:
    python url_scraper.py urls.xlsx             # Excel input
    python url_scraper.py urls.csv              # CSV input
    python url_scraper.py urls.xlsx -o out.docx # Custom output name

Excel/CSV format:
    - Must have a column named 'URL' or 'url' (or the first column is used)
    - One URL per row
"""

import sys
import os
import time
import argparse
import requests
from bs4 import BeautifulSoup
from bs4.element import NavigableString, Tag
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
import csv
from datetime import datetime
from urllib.parse import urljoin


# ── Config ──────────────────────────────────────────────────────────────────

HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/120.0.0.0 Safari/537.36"
    )
}
REQUEST_TIMEOUT = 15       # seconds (default)
DELAY_BETWEEN_REQUESTS = 1 # seconds — be polite
MAX_TEXT_LENGTH = 0      # characters per URL (0 = unlimited, default)


def _url_candidates(url: str) -> list[str]:
    url = url.strip()
    if not url:
        return []
    if url.startswith(("http://", "https://")):
        return [url]
    return [f"https://{url}", f"http://{url}"]


def _add_hyperlink(paragraph, text: str, url: str):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), paragraph.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True))

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    new_run.append(r_pr)

    text_element = OxmlElement("w:t")
    text_element.text = text
    new_run.append(text_element)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def _extract_paragraph_blocks(main: Tag, base_url: str) -> list[list[dict[str, str | None]]]:
    blocks: list[list[dict[str, str | None]]] = []

    for elem in main.find_all(["p", "h1", "h2", "h3", "h4", "li"]):
        segments: list[dict[str, str | None]] = []
        plain_parts: list[str] = []

        def push_text(text: str):
            cleaned = " ".join(text.split())
            if cleaned:
                segments.append({"text": cleaned, "href": None})
                plain_parts.append(cleaned)

        for child in elem.children:
            if isinstance(child, NavigableString):
                push_text(str(child))
            elif isinstance(child, Tag) and child.name == "a":
                link_text = " ".join(child.get_text(" ", strip=True).split())
                if not link_text:
                    continue
                href = child.get("href", "").strip()
                if href:
                    href = urljoin(base_url, href)
                visible_text = link_text if not href else f"{link_text} ({href})"
                segments.append({"text": link_text, "href": href or None})
                plain_parts.append(visible_text)
            elif isinstance(child, Tag):
                push_text(child.get_text(" ", strip=True))

        plain_text = " ".join(" ".join(plain_parts).split())
        if plain_text and len(plain_text) > 20:
            blocks.append(segments)

    return blocks


# ── Input readers ────────────────────────────────────────────────────────────

def read_urls_from_excel(path: str) -> list[str]:
    wb = openpyxl.load_workbook(path)
    ws = wb.active
    headers = [str(c.value).strip().lower() if c.value else "" for c in next(ws.iter_rows(min_row=1, max_row=1))]
    
    # Find the URL column
    url_col = None
    for i, h in enumerate(headers):
        if h in ("url", "urls", "link", "links", "website"):
            url_col = i
            break
    if url_col is None:
        url_col = 0  # fallback: first column
        start_row = 1
    else:
        start_row = 2  # skip header row

    urls = []
    for row in ws.iter_rows(min_row=start_row, values_only=True):
        val = row[url_col] if len(row) > url_col else None
        if val:
            url = str(val).strip()
            if url:
                urls.append(url)
    return urls


def read_urls_from_csv(path: str) -> list[str]:
    urls = []
    with open(path, newline="", encoding="utf-8-sig") as f:
        reader = csv.reader(f)
        headers = next(reader, [])
        url_col = None
        for i, h in enumerate(headers):
            if h.strip().lower() in ("url", "urls", "link", "links", "website"):
                url_col = i
                break
        if url_col is None:
            url_col = 0
            # If first row had no header matching, treat it as a URL
            if headers:
                urls.append(headers[0].strip())

        for row in reader:
            val = row[url_col].strip() if len(row) > url_col else ""
            if val:
                urls.append(val)
    return urls


def read_urls(path: str) -> list[str]:
    ext = os.path.splitext(path)[1].lower()
    if ext in (".xlsx", ".xls"):
        return read_urls_from_excel(path)
    elif ext == ".csv":
        return read_urls_from_csv(path)
    else:
        # Try plain text — one URL per line
        with open(path, encoding="utf-8") as f:
            return [line.strip() for line in f if line.strip()]


# ── Scraper ──────────────────────────────────────────────────────────────────

def scrape_url(url: str, timeout: int = None, max_text: int = None) -> tuple[str, str, str, str, list[list[dict[str, str | None]]]]:
    """
    Returns (resolved_url, title, text, error_message, content_blocks).
    error_message is empty string on success.
    timeout and max_text override the module-level defaults when provided.
    """
    _timeout  = timeout  if timeout  is not None else REQUEST_TIMEOUT
    _max_text = max_text if max_text is not None else MAX_TEXT_LENGTH
    candidates = _url_candidates(url)
    last_error = ""

    for candidate in candidates:
        try:
            resp = requests.get(candidate, headers=HEADERS, timeout=_timeout)
            resp.raise_for_status()
            soup = BeautifulSoup(resp.text, "html.parser")

            title_tag = soup.find("title")
            title = title_tag.get_text(strip=True) if title_tag else "(No title)"

            for tag in soup(["script", "style", "nav", "footer", "header",
                              "aside", "noscript", "form", "button", "svg", "img"]):
                tag.decompose()

            main = (soup.find("main") or soup.find("article") or
                    soup.find(id="content") or soup.find(class_="content") or
                    soup.find("body") or soup)

            blocks = _extract_paragraph_blocks(main, candidate)
            paragraphs = []
            for block in blocks:
                text = " ".join(segment["text"] for segment in block if segment.get("text"))
                if text:
                    paragraphs.append(text)

            text = "\n\n".join(paragraphs)

            if _max_text and len(text) > _max_text:
                text = text[:_max_text] + "\n\n[... content truncated ...]"

            if not text.strip():
                text = "(No readable text content found on this page.)"

            return candidate, title, text, "", blocks

        except requests.exceptions.Timeout:
            last_error = f"Request timed out after {_timeout}s"
        except requests.exceptions.ConnectionError as e:
            last_error = f"Connection error: {e}"
        except requests.exceptions.HTTPError as e:
            last_error = f"HTTP {e.response.status_code}: {e}"
        except Exception as e:
            last_error = f"Unexpected error: {e}"

    return (candidates[0] if candidates else url), "", "", last_error or "Unexpected error", []


# ── Word doc builder ─────────────────────────────────────────────────────────

def build_docx(results: list[dict], output_path: str) -> None:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── Cover title ────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("Web Content Scrape Report")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run(
        f"Generated: {datetime.now().strftime('%B %d, %Y %H:%M')}  ·  "
        f"{len(results)} URL(s) processed"
    )
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()  # spacer

    # ── Each URL ──────────────────────────────────────────────────────────
    for i, r in enumerate(results, 1):
        # URL heading
        url_heading = doc.add_paragraph()
        idx_run = url_heading.add_run(f"URL {i}: ")
        idx_run.bold = True
        idx_run.font.size = Pt(13)
        idx_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

        url_run = url_heading.add_run(r["url"])
        url_run.bold = True
        url_run.font.size = Pt(13)
        url_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        url_run.underline = True

        # Page title (if available)
        if r.get("title"):
            tp = doc.add_paragraph()
            t = tp.add_run(f"Page Title: {r['title']}")
            t.italic = True
            t.font.size = Pt(10)
            t.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        # Error or content
        if r["error"]:
            ep = doc.add_paragraph()
            er = ep.add_run(f"⚠ Could not scrape: {r['error']}")
            er.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            er.font.size = Pt(11)
        else:
            # Content paragraphs
            blocks = r.get("content_blocks") or []
            if blocks:
                for block in blocks:
                    cp = doc.add_paragraph()
                    for segment in block:
                        text = (segment.get("text") or "").strip()
                        href = segment.get("href")
                        if not text:
                            continue
                        if href:
                            _add_hyperlink(cp, text, href)
                        else:
                            cp.add_run(text).font.size = Pt(11)
            else:
                for line in r["text"].split("\n\n"):
                    line = line.strip()
                    if not line:
                        continue
                    cp = doc.add_paragraph()
                    cr = cp.add_run(line)
                    cr.font.size = Pt(11)

        # Divider
        doc.add_paragraph("─" * 80)
        doc.add_paragraph()

    doc.save(output_path)


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Scrape URLs and save to Word doc.")
    parser.add_argument("input", help="Excel (.xlsx) or CSV file with URLs")
    parser.add_argument("-o", "--output", default="scraped_content.docx",
                        help="Output .docx filename (default: scraped_content.docx)")
    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"❌ File not found: {args.input}")
        sys.exit(1)

    print(f"📂 Reading URLs from: {args.input}")
    urls = read_urls(args.input)

    if not urls:
        print("❌ No URLs found in the file.")
        sys.exit(1)

    print(f"🔗 Found {len(urls)} URL(s)\n")

    results = []
    for i, url in enumerate(urls, 1):
        print(f"[{i}/{len(urls)}] Scraping: {url}")
        resolved_url, title, text, error, content_blocks = scrape_url(url)

        if error:
            print(f"    ⚠  {error}")
        else:
            words = len(text.split())
            print(f"    ✅ Title: {title[:60]}  |  ~{words} words")

        results.append({"url": resolved_url, "title": title, "text": text, "error": error, "content_blocks": content_blocks})

        if i < len(urls):
            time.sleep(DELAY_BETWEEN_REQUESTS)

    print(f"\n💾 Saving to: {args.output}")
    build_docx(results, args.output)
    print(f"✅ Done! File saved: {args.output}")


if __name__ == "__main__":
    main()
